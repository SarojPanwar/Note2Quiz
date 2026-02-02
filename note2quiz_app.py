import streamlit as st
import re
import time
import os
import io
import pandas as pd 
from dotenv import load_dotenv 
import requests  
from docx import Document   

load_dotenv() 

os.environ["GOOGLE_API_KEY"]=os.getenv("GEMINI_API_KEY")
    
STRICT_SYSTEM_INSTRUCTION = (
     "You are a highly analytical academic question generator. "
    "Your only task is to create exam-based questions from the given text. "
    "Avoid conversational language, greetings, or explanations. "
    "Return only a clean, numbered list of short, exam-style questions."
)



def classify_bloom(question):
    """Classifies a question based on simple keyword matching (Bloom's Taxonomy)."""
    q = question.lower()
    if any(x in q for x in ["define", "list", "name", "what is", "who is"]):
        return "Knowledge"
    elif any(x in q for x in ["explain", "summarize", "describe", "identify"]):
        return "Comprehension"
    elif any(x in q for x in ["apply", "use", "solve", "demonstrate"]):
        return "Application"
    elif any(x in q for x in ["analyze", "compare", "contrast", "why", "examine"]):
        return "Analysis"
    elif any(x in q for x in ["design", "compose", "create", "what if", "develop"]):
        return "Synthesis"
    elif any(x in q for x in ["evaluate", "assess", "argue", "critique", "justify"]):
        return "Evaluation"
    else:
        return "Unclassified"
    
def extract_text_from_pdf(uploaded_file):
    """Extracts clean, properly spaced text from a PDF using PyMuPDF."""

    import fitz  # PyMuPDF
   
    text = ""
    try:
       
        with fitz.open(stream=uploaded_file.read(), filetype="pdf") as pdf:
            for page in pdf:
                page_text = page.get_text("text")  
                if page_text:
                    text += " " + page_text

       
        text = " ".join(text.split()) 
        text = (
            text.replace(" ,", ",")
                .replace(" .", ".")
                .replace(" :", ":")
                .replace(" ;", ";")
                .replace(" ?", "?")
                .replace(" !", "!")
        )

    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return ""

    return text.strip()

def extract_text_from_csv(uploaded_file):
    """Extracts and cleans text from a CSV file using pandas."""
    try:
        uploaded_file.seek(0)
        df = pd.read_csv(uploaded_file)

        
        text = df.to_string(index=False, header=True)

      
        text = " ".join(text.split()) 
        text = (
            text.replace(" ,", ",")
                .replace(" .", ".")
                .replace(" :", ":")
                .replace(" ;", ";")
        )

        return text.strip()

    except Exception as e:
        st.error(f"Error reading CSV file: {e}")
        return ""        
    



MODEL_NAME = "gemini-2.5-flash"
def call_gemini(prompt):
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        st.error("GEMINI_API_KEY not found")
        return ""

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL_NAME}:generateContent"
    headers = {"Content-Type": "application/json"}
    params = {"key": api_key}

    payload = {
        "contents": [
            {
                "parts": [
                    {"text": prompt}
                ]
            }
        ]
    }
    time.sleep(2)
    try:
        r = requests.post(
            url,
            headers=headers,
            params=params,
            json=payload,
            timeout=60
        )
        if r.status_code == 429:
            st.warning("AI is busy right now. Please wait a few seconds ⏳")
            return ""
        r.raise_for_status()
        return r.json()["candidates"][0]["content"]["parts"][0]["text"]

    except requests.exceptions.RequestException:
        st.error("Network or API issue. Please try again later.")
        return ""







def generate_mcqs(text, num_questions=10):
    """Generates exam-based short MCQs using the Gemini API."""

    prompt = STRICT_SYSTEM_INSTRUCTION + "\n\n" + f"""
You are an expert exam question setter.

Generate {num_questions} **short, exam-style multiple-choice questions (MCQs)** 
based strictly on the following notes. 

Each MCQ must:
- Be concise and relevant (max 1 sentence).
- Have **4 options (A, B, C, D)**.
- Have exactly **one correct answer**.
- Be suitable for college-level objective exams.

Provide output in this exact format:
1. Question?
A) Option 1
B) Option 2
C) Option 3
D) Option 4
Answer: C) Correct Option

Notes:
{text[:4000]}
    """

    try:
        raw_text = call_gemini(prompt)
        if not raw_text:
            return []

        raw = raw_text.strip().split("\n")

        questions, current_q = [], {}

        for line in raw:
            line = line.strip()
            if line.startswith(tuple(str(i)+'.' for i in range(1, 100))):
                if current_q:
                    questions.append(current_q)
                    current_q = {}
                current_q["question"] = line
                current_q["options"] = []
            elif line.startswith(("A)", "B)", "C)", "D)")):
                current_q.setdefault("options", []).append(line)
            elif line.lower().startswith("answer:"):
                current_q["answer"] = line.split(":", 1)[-1].strip()
            if len(questions) >= num_questions:
                break

        if current_q:
            questions.append(current_q)

        formatted = []
        for q in questions[:num_questions]:
            formatted.append({
                "question": q.get("question", ""),
                "options": [(opt[3:].strip() if len(opt) > 2 else opt.strip()) 
            for opt in q.get("options", [])],
              
                "answer": q.get("answer", ""),
                "bloom": classify_bloom(q.get("question", "")),
            })
        return formatted

    except Exception as e:
        st.error(f"Gemini API Call Failed (MCQs): {e}")
        return []
    
def generate_viva_questions(text, num_questions=5):

    prompt = f"""
    You are an academic examiner.
    Based on the following lecture notes, generate exactly {num_questions} high-level viva questions (Analysis, Synthesis, or Evaluation).
    Rules:
    -Output ONLY numbered questions
    -one question per line
    -No explanationa or extra text
    Notes:
    ---
    {text[:4000]} 
    """

    try:
        raw_text = call_gemini(prompt)
        if not raw_text:
            return []
       
        matches = re.findall(r"\d+\.\s*(.+)",raw_text)
        
        questions = []
        for q  in matches[:num_questions]:
            questions.append({
                "question":q.strip(),
                "bloom" :classify_bloom(q)
            })
        return questions
        
    except Exception as e:
        st.error(f"Gemini API Call Failed (Viva Questions): {e}")
        return []   

def create_word_document(mcqs, viva_questions):
    """Creates a .docx file from the generated questions."""
    document = Document()
    document.add_heading('Note2Quiz Generated Questions', 0)
    
   
    document.add_heading('Multiple Choice Questions (MCQs)', level=1)
    for i, q in enumerate(mcqs, 1):
        document.add_paragraph(f"Q{i}: {q['question']}", style='List Number')
        for idx, opt in enumerate(q['options']):
            document.add_paragraph(f"  {chr(65+idx)}. {opt}", style='List Bullet')
        document.add_paragraph(f"✅ Answer: {q['answer']} | 🧠 Bloom Level: {q['bloom']}").bold = True
        document.add_paragraph("\n")

    
    document.add_heading('Viva/Discussion Questions', level=1)
    for i, q in enumerate(viva_questions, 1):
        document.add_paragraph(f"Q{i}: {q['question']}", style='List Number')
        document.add_paragraph(f"🧠 Bloom Level: {q['bloom']}").bold = True
        document.add_paragraph("\n") 
    
    
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io   


def main():
  
    st.set_page_config(
        page_title="Note2Quiz - Smart Question Generator",
        page_icon="🧠",
        layout="wide",
        initial_sidebar_state="expanded"
    )

  
    st.sidebar.title("⚙️ Settings")
    num_mcqs = st.sidebar.slider("Number of MCQs", 5, 20, 10)
    num_viva = st.sidebar.slider("Number of Viva Questions", 3, 10, 5)
    st.sidebar.markdown("---")
    st.sidebar.info("Upload your notes file and generate questions instantly.")

    st.markdown("""
    <div style="text-align:center; padding: 10px 0;">
        <h1 style="color:#0066FF;">🧠 Note2Quiz</h1>
        <h3>Transform Lecture Notes into Smart Exam Questions</h3>
        <p>Upload your PDF or CSV file to generate MCQs and Viva Questions tagged with Bloom’s Taxonomy.</p>
    </div>
    """, unsafe_allow_html=True)

   
    tab1, tab2 = st.tabs(["📤 Upload File", "📊 Generated Questions"])

    with tab1:
        st.header("Step 1: Upload Your File")
        uploaded_file = st.file_uploader("📄 Upload a PDF or CSV file", type=["pdf", "csv"])

        if uploaded_file:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            text = ""

            if file_extension == "pdf":
                with st.spinner("Extracting text from PDF..."):
                    text = extract_text_from_pdf(uploaded_file)
            elif file_extension == "csv":
                with st.spinner("Extracting text from CSV..."):
                    text = extract_text_from_csv(uploaded_file)
            else:
                st.error("❌ Invalid file format. Please provide a PDF or CSV file.")
                st.stop()

            if not text:
                st.warning("⚠️ Could not extract meaningful text. Try another file.")
                st.stop()
            st.session_state["text"] =text

           
            with st.spinner(f"Generating {num_mcqs} MCQs and {num_viva} Viva Questions..."):
                mcqs = generate_mcqs(text, num_questions=num_mcqs)
                viva = generate_viva_questions(text, num_questions=num_viva)

           
            st.session_state["results"] = (mcqs, viva)
            st.success("✅ Questions generated successfully! Go to the 'Generated Questions' tab to view them.") 

    with tab2:
      st.header("Step 2: View Generated Questions")

      if "results" in st.session_state:
          mcqs, viva = st.session_state["results"]
      
          with st.expander("📘 View Generated MCQs"):
                if mcqs:
                  for i, q in enumerate(mcqs, 1):
                   
                    st.markdown(f"**Q{i}:** {q['question']}")
                   
                    for idx, opt in enumerate(q['options']):
                       
                        st.markdown(f"**{chr(65+idx)}.** {opt}")
                        
                    
                    st.markdown("---")
                    st.markdown(
                       
                        f"**✅ Correct Answer:** {q['answer']} | **🧠 Bloom Level:** **{q['bloom']}**"
                    )
                    st.markdown("---") 
                else:
                  st.info("No MCQs were generated. Please check the content or try regenerating.") 
                
        
        
          with st.expander("🎤 View Viva Questions"):
              if viva:
                  for i, q in enumerate(viva, 1):
                    st.markdown(f"**Q{i}:** {q['question']}")
                    st.markdown(f"🧠 **Bloom Level:** **{q['bloom']}**")
                    st.markdown("---")
              else:
                 st.info("No Viva Questions were generated.")

      
          st.markdown("---")
          st.subheader("💾 Export Options")
        
          if mcqs or viva:
           
              doc_io = create_word_document(mcqs, viva) 
              st.download_button(
                label="Download as Word Document (.docx)",
                data=doc_io,
                file_name="Note2Quiz_Questions.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
          else:
             st.warning("Cannot export. Please generate questions first.")

      else:
          st.info("ℹ️ Upload a file first in the 'Upload File' tab.")

      if "text" in st.session_state:
          
        user_input = st.chat_input("Ask a question from your uploaded notes...")
        response=None
        if user_input:
         prompt =f"""
         You are a teaching assistant.
         Answer the question ONLY using the following notes.
         If the answer is not present ,say "Answer not found in the notes"
         Notes:
         {st.session_state["text"][:3000]}
         Question:
         {user_input}   
         """
         with st.spinner("Thinking..."):
             response = call_gemini(prompt)
        if response:
            st.markdown("### 🤖 Answer")
            st.write(response)
             
             
            
            
  
    st.markdown("""
    <hr>
    <div style="text-align:center; color:grey;">
      #  <p>Developed with ❤️| Powered by Streamlit & Gemini</p>
    </div>
    """, unsafe_allow_html=True)
if __name__ == "__main__":
    main()    
